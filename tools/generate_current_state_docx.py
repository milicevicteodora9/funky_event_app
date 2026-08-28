from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "documents" / "trenutno-stanje-aplikacije.docx"


def set_font(run, size, bold=False, color="17324D"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_bottom_border(paragraph, color="2E6F95", size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("FUNKY EVENT APP  |  PREGLED PROJEKTA")
    set_font(hr, 8.5, bold=True, color="6B7785")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(22)
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("Trenutno stanje aplikacije")
    set_font(tr, 24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    sr = subtitle.add_run("Kratak pregled faze razvoja")
    set_font(sr, 12, color="5F6B76")
    add_bottom_border(subtitle)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run(
        "Funky Event App se trenutno nalazi u fazi razvoja korisničkog interfejsa i projektnog modela sistema. "
        "Implementirani ekrani prikazuju planirane funkcionalnosti za upravljanje klijentima, događajima, budžetima, "
        "blagajnom, timom, računima i fakturama."
    )

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.add_run(
        "Za prikaz i testiranje interfejsa trenutno se koriste mock podaci iz klase MockDataRepository, zbog čega aplikacija "
        "još nema trajno čuvanje podataka. Backend, baza podataka i povezivanje sa stvarnim servisima biće realizovani u narednoj "
        "fazi razvoja, dok postojeća verzija predstavlja funkcionalan UI prototip i osnovu za dalju implementaciju."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Funky Event App - UI prototip")
    set_font(fr, 8.5, color="7A8793")

    doc.core_properties.title = "Trenutno stanje aplikacije"
    doc.core_properties.subject = "Funky Event App - UI prototip i planirana backend faza"
    doc.core_properties.author = "Funky Event App"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
